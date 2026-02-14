"""
Document Processor Module
Handles loading, parsing, and chunking documents of various formats.
Supports: PDF, DOCX, TXT, MD, CSV
"""

import os
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime

import nltk
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)


# ──────────────────────────────────────────────
# Data Models
# ──────────────────────────────────────────────

@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document."""
    content: str
    metadata: Dict = field(default_factory=dict)
    chunk_id: str = ""
    doc_id: str = ""
    chunk_index: int = 0

    def __post_init__(self):
        if not self.chunk_id:
            self.chunk_id = hashlib.md5(
                f"{self.doc_id}_{self.chunk_index}_{self.content[:50]}".encode()
            ).hexdigest()


@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""
    doc_id: str
    filename: str
    file_type: str
    raw_text: str
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)
    processed_at: str = ""

    def __post_init__(self):
        if not self.processed_at:
            self.processed_at = datetime.now().isoformat()
        if not self.doc_id:
            self.doc_id = hashlib.md5(
                f"{self.filename}_{self.raw_text[:100]}".encode()
            ).hexdigest()


# ──────────────────────────────────────────────
# Document Loaders
# ──────────────────────────────────────────────

class PDFLoader:
    """Extract text from PDF files."""

    @staticmethod
    def load(file_path: str) -> Tuple[str, Dict]:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(text)

        metadata = {
            "total_pages": len(reader.pages),
            "pdf_info": {k: str(v) for k, v in reader.metadata.items()} if reader.metadata else {},
        }
        return "\n\n".join(pages), metadata


class DOCXLoader:
    """Extract text from DOCX files."""

    @staticmethod
    def load(file_path: str) -> Tuple[str, Dict]:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_texts.append(" | ".join(row_data))

        full_text = "\n\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n[TABLE DATA]\n" + "\n".join(table_texts)

        metadata = {
            "total_paragraphs": len(paragraphs),
            "total_tables": len(doc.tables),
        }
        return full_text, metadata


class TextLoader:
    """Extract text from plain text and markdown files."""

    @staticmethod
    def load(file_path: str) -> Tuple[str, Dict]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        metadata = {
            "total_lines": text.count("\n") + 1,
            "total_chars": len(text),
        }
        return text, metadata


class CSVLoader:
    """Extract text from CSV files."""

    @staticmethod
    def load(file_path: str) -> Tuple[str, Dict]:
        import csv

        rows = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            headers = next(reader, [])
            for row in reader:
                row_dict = {h: v for h, v in zip(headers, row)}
                rows.append(" | ".join(f"{k}: {v}" for k, v in row_dict.items()))

        text = f"Headers: {', '.join(headers)}\n\n" + "\n".join(rows)
        metadata = {
            "total_rows": len(rows),
            "headers": headers,
        }
        return text, metadata


# ──────────────────────────────────────────────
# Text Chunking Strategies
# ──────────────────────────────────────────────

class TextChunker:
    """
    Intelligent text chunking with multiple strategies.
    """

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        strategy: str = "semantic",  # "fixed", "sentence", "semantic", "paragraph"
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.strategy = strategy

    def chunk(self, text: str) -> List[str]:
        """Chunk text using the configured strategy."""
        strategies = {
            "fixed": self._fixed_chunk,
            "sentence": self._sentence_chunk,
            "semantic": self._semantic_chunk,
            "paragraph": self._paragraph_chunk,
        }
        chunker = strategies.get(self.strategy, self._semantic_chunk)
        chunks = chunker(text)

        # Filter out empty or very short chunks
        return [c.strip() for c in chunks if len(c.strip()) > 20]

    def _fixed_chunk(self, text: str) -> List[str]:
        """Simple fixed-size character chunking with overlap."""
        chunks = []
        start = 0
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - self.chunk_overlap
        return chunks

    def _sentence_chunk(self, text: str) -> List[str]:
        """Chunk by sentences, respecting chunk_size limits."""
        sentences = nltk.sent_tokenize(text)
        chunks = []
        current_chunk = []
        current_length = 0

        for sentence in sentences:
            sentence_len = len(sentence)
            if current_length + sentence_len > self.chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                # Keep overlap sentences
                overlap_text = ""
                overlap_sents = []
                for s in reversed(current_chunk):
                    if len(overlap_text) + len(s) <= self.chunk_overlap:
                        overlap_sents.insert(0, s)
                        overlap_text += s
                    else:
                        break
                current_chunk = overlap_sents
                current_length = len(overlap_text)

            current_chunk.append(sentence)
            current_length += sentence_len

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return chunks

    def _semantic_chunk(self, text: str) -> List[str]:
        """
        Semantic chunking: split by sections/headers first,
        then fall back to sentence chunking within each section.
        """
        # Detect section boundaries (markdown headers, numbered sections, etc.)
        section_pattern = r'(?:^|\n)(?:#{1,6}\s+.+|(?:\d+\.)+\s+.+|[A-Z][A-Z\s]{5,}(?:\n|$))'
        sections = re.split(section_pattern, text)

        if len(sections) <= 1:
            # No clear sections found — fall back to sentence chunking
            return self._sentence_chunk(text)

        # Chunk within each section
        all_chunks = []
        for section in sections:
            section = section.strip()
            if not section:
                continue
            if len(section) <= self.chunk_size:
                all_chunks.append(section)
            else:
                sub_chunks = self._sentence_chunk(section)
                all_chunks.extend(sub_chunks)

        return all_chunks

    def _paragraph_chunk(self, text: str) -> List[str]:
        """Chunk by paragraphs, merging small ones."""
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            para_len = len(para)

            if current_length + para_len > self.chunk_size and current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_length = 0

            current_chunk.append(para)
            current_length += para_len

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return chunks


# ──────────────────────────────────────────────
# Main Document Processor
# ──────────────────────────────────────────────

class DocumentProcessor:
    """
    Main class for processing documents end-to-end.
    Loads → Extracts text → Chunks → Returns ProcessedDocument
    """

    LOADERS = {
        ".pdf": PDFLoader,
        ".docx": DOCXLoader,
        ".doc": DOCXLoader,
        ".txt": TextLoader,
        ".md": TextLoader,
        ".csv": CSVLoader,
    }

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        chunking_strategy: str = "semantic",
    ):
        self.chunker = TextChunker(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            strategy=chunking_strategy,
        )

    def process(self, file_path: str) -> ProcessedDocument:
        """Process a single document file."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.LOADERS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {list(self.LOADERS.keys())}"
            )

        # Load document
        loader = self.LOADERS[ext]
        raw_text, doc_metadata = loader.load(str(path))

        if not raw_text.strip():
            raise ValueError(f"No text could be extracted from: {file_path}")

        # Create document
        doc = ProcessedDocument(
            doc_id="",
            filename=path.name,
            file_type=ext,
            raw_text=raw_text,
            metadata={
                "file_path": str(path.absolute()),
                "file_size_bytes": path.stat().st_size,
                **doc_metadata,
            },
        )

        # Chunk the text
        chunk_texts = self.chunker.chunk(raw_text)
        doc.chunks = [
            DocumentChunk(
                content=chunk_text,
                metadata={
                    "source": path.name,
                    "chunk_index": i,
                    "total_chunks": len(chunk_texts),
                    **doc_metadata,
                },
                doc_id=doc.doc_id,
                chunk_index=i,
            )
            for i, chunk_text in enumerate(chunk_texts)
        ]

        return doc

    def process_directory(self, dir_path: str) -> List[ProcessedDocument]:
        """Process all supported documents in a directory."""
        docs = []
        directory = Path(dir_path)

        for ext in self.LOADERS:
            for file_path in directory.glob(f"*{ext}"):
                try:
                    doc = self.process(str(file_path))
                    docs.append(doc)
                    print(f"  ✓ Processed: {file_path.name} → {len(doc.chunks)} chunks")
                except Exception as e:
                    print(f"  ✗ Failed: {file_path.name} → {e}")

        return docs
